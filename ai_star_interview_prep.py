import io
import os
import html
import re
import importlib

import streamlit as st
from huggingface_hub import InferenceClient # type: ignore
from docx import Document
from docx.shared import Pt


st.set_page_config(page_title="AI-STAR Interview Prep", layout="wide")

st.markdown(
    """
    <style>
    /* Global bold text styling */
    .stApp,
    .stApp p,
    .stApp span,
    .stApp label,
    .stApp div,
    .stApp li,
    .stApp h1,
    .stApp h2,
    .stApp h3,
    .stApp h4,
    .stApp h5,
    .stApp h6 {
        font-weight: 700 !important;
    }

    /* Keep main title/subheaders as-is, but make all other UI text 16px */
    .stApp p,
    .stApp span,
    .stApp label,
    .stApp li,
    .stApp small,
    .stButton > button,
    .stDownloadButton > button,
    .stCheckbox label {
        font-size: 16px !important;
    }

    div[data-testid="stTextInput"] input,
    div[data-testid="stTextArea"] textarea {
        font-size: 16px !important;
    }

    div[data-testid="stTextInput"] input::placeholder,
    div[data-testid="stTextArea"] textarea::placeholder {
        font-size: 16px !important;
    }

    /* Hide Streamlit helper text like 'Press Enter to apply' under inputs */
    [data-testid="InputInstructions"] {
        display: none !important;
    }

    /* Professional neutral styling for top behavioral question input */
    div[data-testid="stTextInput"] input {
        width: 100% !important;
        background: #f8fafc !important;
        color: #0f172a !important;
        border: 1px solid #cbd5e1 !important;
        border-radius: 10px !important;
    }

    div[data-testid="stTextInput"] input::placeholder {
        color: #64748b !important;
    }

    div[data-testid="stTextInput"] input:focus {
        border: 1px solid #2563eb !important;
        box-shadow: 0 0 0 2px rgba(37, 99, 235, 0.2) !important;
    }

    /* Professional neutral styling for STAR narrative text areas */
    div[data-testid="stTextArea"] textarea {
        width: 100% !important;
        background: #f8fafc !important;
        color: #0f172a !important;
        border: 1px solid #cbd5e1 !important;
        border-radius: 10px !important;
    }

    div[data-testid="stTextArea"] textarea::placeholder {
        color: #64748b !important;
    }

    div[data-testid="stTextArea"] textarea:focus {
        border: 1px solid #2563eb !important;
        box-shadow: 0 0 0 2px rgba(37, 99, 235, 0.2) !important;
    }

    div[data-testid="stTextArea"] label {
        color: #111827 !important;
        font-weight: 600 !important;
    }

    /* Green full-width download button */
    div.stDownloadButton > button {
        width: 100% !important;
        background: #16a34a !important;
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        font-weight: 700 !important;
        padding: 0.65rem 1rem !important;
    }

    div.stDownloadButton > button:hover {
        background: #15803d !important;
    }

    .answer-shell {
        min-height: 360px;
        border: 1px solid #e5e7eb;
        border-radius: 12px;
        background: #ffffff;
        padding: 1rem;
    }

    .star-section {
        margin-bottom: 0.85rem;
    }

    .star-header {
        font-weight: 700 !important;
        color: #0f172a;
        margin-bottom: 0.2rem;
    }

    .answer-shell .star-body,
    .answer-shell .star-body p,
    .answer-shell .star-body span,
    .answer-shell .star-body div,
    .answer-shell .star-body li,
    .answer-shell .star-body ol,
    .answer-shell .star-body ul {
        font-weight: 400 !important;
        color: #1f2937;
        line-height: 1.5;
    }

    .star-list {
        margin: 0.25rem 0 0.25rem 1.25rem;
        padding-left: 0.6rem;
    }

    .answer-shell .star-list li::marker {
        font-weight: 700 !important;
        color: #0f172a;
    }

    .brand-row {
        display: flex;
        align-items: center;
        gap: 12px;
        margin-bottom: 0.5rem;
    }

    .star-logo {
        width: 44px;
        height: 44px;
        border-radius: 10px;
        background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
        color: #fbbf24;
        border: 1px solid #f59e0b;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 24px;
        font-weight: 700;
        line-height: 1;
        box-shadow: 0 4px 12px rgba(37, 99, 235, 0.25);
    }

    .brand-title {
        margin: 0;
        color: #0f172a;
        font-size: 2rem;
        line-height: 1.15;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


HF_MODEL = os.getenv("HF_MODEL", "microsoft/Phi-3-mini-4k-instruct")
FALLBACK_MODELS = [
    "meta-llama/Llama-3.1-8B-Instruct",
    "Qwen/Qwen2.5-7B-Instruct",
    "mistralai/Mistral-7B-Instruct-v0.3",
]

# Broad profanity checks with separator-tolerant patterns (e.g., f-u-c-k, f.u.c.k).
PROFANITY_PATTERNS = [
    r"\bf[\W_]*u[\W_]*c[\W_]*k(?:[\W_]*e[\W_]*r|[\W_]*i[\W_]*n[\W_]*g|[\W_]*s)?\b",
    r"\bs[\W_]*h[\W_]*i[\W_]*t(?:[\W_]*t[\W_]*y|[\W_]*s)?\b",
    r"\bb[\W_]*i[\W_]*t[\W_]*c[\W_]*h(?:[\W_]*e[\W_]*s|[\W_]*y)?\b",
    r"\ba[\W_]*s[\W_]*s(?:[\W_]*h[\W_]*o[\W_]*l[\W_]*e|[\W_]*e[\W_]*s)?\b",
    r"\bd[\W_]*a[\W_]*m[\W_]*n(?:[\W_]*e[\W_]*d|[\W_]*i[\W_]*t)?\b",
    r"\bc[\W_]*r[\W_]*a[\W_]*p(?:[\W_]*p[\W_]*y)?\b",
    r"\bp[\W_]*i[\W_]*s[\W_]*s(?:[\W_]*e[\W_]*d|[\W_]*i[\W_]*n[\W_]*g)?\b",
    r"\bd[\W_]*i[\W_]*c[\W_]*k(?:[\W_]*h[\W_]*e[\W_]*a[\W_]*d)?\b",
    r"\bb[\W_]*a[\W_]*s[\W_]*t[\W_]*a[\W_]*r[\W_]*d\b",
    r"\bm[\W_]*o[\W_]*t[\W_]*h[\W_]*e[\W_]*r[\W_]*f[\W_]*u[\W_]*c[\W_]*k[\W_]*e[\W_]*r\b",
]
PROFANITY_REGEXES = [re.compile(pattern, re.IGNORECASE) for pattern in PROFANITY_PATTERNS]



def get_hf_token() -> str:
    # Local runs may not have a secrets.toml; accessing st.secrets can raise.
    def _safe_secret_get(key: str) -> str:
        try:
            return st.secrets.get(key, "")
        except Exception:
            return ""

    token = os.getenv("HF_TOKEN", "")
    if not token:
        token = _safe_secret_get("HF_TOKEN")
    if not token:
        token = _safe_secret_get("HF_API_TOKEN")
    if not token:
        token = os.getenv("HF_API_TOKEN", "")
    if not token:
        token = os.getenv("HUGGINGFACEHUB_API_TOKEN", "")
    return token


def has_profanity(text: str) -> bool:
    if not text or not text.strip():
        return False
    return any(regex.search(text) for regex in PROFANITY_REGEXES)


def any_profanity_in_fields(fields: list[str]) -> bool:
    return any(has_profanity(field) for field in fields)


@st.cache_resource
def get_inference_client(token: str) -> InferenceClient:
    return InferenceClient(api_key=token)


def _candidate_models() -> list[str]:
    # Keep preferred model first, then try known fallbacks if provider support differs.
    models = [HF_MODEL] + FALLBACK_MODELS
    deduped = []
    for model in models:
        if model and model not in deduped:
            deduped.append(model)
    return deduped


def build_prompt(question: str, situation: str, task: str, action: str, result: str, add_followups: bool) -> str:
    followup_instruction = (
        "Then add exactly 2 tailored follow-up interview questions after the final answer under the exact header 'Follow-up Questions'. "
        "Output them as 2 separate numbered lines only (1. and 2.) with no extra explanation."
        if add_followups
        else "Do not add follow-up questions."
    )

    return f"""You are an expert interview coach.
Create a polished, professional STAR interview response in first person.

Behavioral Question:
{question}

Candidate STAR Notes:
Situation: {situation}
Task: {task}
Action: {action}
Result: {result}

Requirements:
1) Produce a cohesive final answer that sounds natural, specific, and impactful.
2) Keep the response concise but substantive.
3) Use this exact output structure and exact header text (including punctuation):
The Situation:
<1 short paragraph>

The Task:
<1 short paragraph>

The Action Plan:
To address this pressing concern, I took the following steps
1. <step>
2. <step>
3. <step>
4. <step>

The Result:
<1 short paragraph>

Follow-up Questions
1. <question>
2. <question>

4) Keep the Action Plan as a numbered list (1., 2., 3., ...), with at least 3 steps.
5) If follow-up questions are not requested, omit the entire Follow-up Questions section.
6) {followup_instruction}

Return only the final response text.
"""


def call_hf_inference(prompt: str, token: str, max_tokens: int = 500) -> str:
    client = get_inference_client(token)
    last_error = None

    for model_name in _candidate_models():
        try:
            completion = client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
                stream=False,
            )

            if completion and completion.choices and completion.choices[0].message:
                content = completion.choices[0].message.content
                if isinstance(content, str) and content.strip():
                    return content.strip()

            last_error = ValueError(f"Unexpected response format for model '{model_name}'.")
        except Exception as e:
            msg = str(e)
            last_error = e
            if "model_not_supported" in msg or "not supported by any provider" in msg:
                continue
            raise

    raise ValueError(
        "No supported model was available for your enabled providers. "
        "Set HF_MODEL in secrets/env to a model available in your HF account/providers. "
        f"Tried: {', '.join(_candidate_models())}. Last error: {last_error}"
    )


def _set_times_new_roman(paragraph, size=12, bold=False):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold


def build_docx_bytes(question: str, situation: str, task: str, action: str, result: str, final_answer: str) -> bytes:
    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    title = doc.add_paragraph("AI-STAR Interview Prep Response")
    _set_times_new_roman(title, size=14, bold=True)
    title.paragraph_format.space_after = Pt(10)

    q_header = doc.add_paragraph("Behavioral Question")
    _set_times_new_roman(q_header, size=12, bold=True)
    q_header.paragraph_format.space_after = Pt(2)

    q_body = doc.add_paragraph(question)
    _set_times_new_roman(q_body, size=12, bold=False)
    q_body.paragraph_format.space_after = Pt(10)

    for header, body in [
        ("Situation", situation),
        ("Task", task),
        ("Action", action),
        ("Result", result),
    ]:
        h = doc.add_paragraph(header)
        _set_times_new_roman(h, size=12, bold=True)
        h.paragraph_format.space_after = Pt(2)

        p = doc.add_paragraph(body)
        _set_times_new_roman(p, size=12, bold=False)
        p.paragraph_format.space_after = Pt(8)

    final_header = doc.add_paragraph("Final Answer")
    _set_times_new_roman(final_header, size=12, bold=True)
    final_header.paragraph_format.space_after = Pt(2)

    final_body = doc.add_paragraph(final_answer)
    _set_times_new_roman(final_body, size=12, bold=False)
    final_body.paragraph_format.space_after = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def extract_resume_text(uploaded_file) -> str:
    if uploaded_file is None:
        return ""

    file_ext = os.path.splitext(uploaded_file.name)[1].lower()

    if file_ext == ".pdf":
        try:
            pdf_module_name = "PyPDF2"
            try:
                pdf_module = importlib.import_module(pdf_module_name)
            except ImportError:
                pdf_module_name = "pypdf"
                pdf_module = importlib.import_module(pdf_module_name)

            PdfReader = getattr(pdf_module, "PdfReader")
            pdf_reader = PdfReader(uploaded_file)
            return " ".join(page.extract_text() or "" for page in pdf_reader.pages).strip()
        except Exception as exc:
            st.error(f"Could not read PDF: {exc}")
            return ""

    if file_ext == ".docx":
        try:
            import docx

            doc = docx.Document(uploaded_file)
            return " ".join(para.text for para in doc.paragraphs).strip()
        except Exception:
            uploaded_file.seek(0)
            return uploaded_file.read().decode("utf-8", errors="ignore").strip()

    uploaded_file.seek(0)
    return uploaded_file.read().decode("utf-8", errors="ignore").strip()


def build_narrative_prompt(project_description: str, resume_text: str) -> str:
    return f"""
You are an expert interview coach.

Using the job description and resume content below, create a "Tell My Story" narrative for a general professional conversation.

Project Description:
{project_description}

Resume Content:
{resume_text}

Instructions:
1. Provide a speaking script for up to 2 minutes at a medium speaking pace.
2. Keep the total length between 210 and 260 words, and never exceed 260 words.
3. Keep it in first person, polished, confident, and natural for live delivery.
4. Use a professional-conversational tone: warm, clear, and human, but not overly casual.
5. Vary sentence length and avoid repetitive sentence openings.
6. Use plain spoken language and light contractions where natural (for example: "I've", "I've led", "I'm excited").
7. Avoid robotic or overly formal phrases like "I am writing to express", "therefore", "moreover", "in conclusion", "it is imperative", "leverage synergies", and "utilize".
8. Sound like a real candidate speaking naturally in an interview, not reading a formal essay.
9. Use my resume experience as the foundation and tightly align it to this specific job description.
10. Intertwine my experience with the role requirements so the response sounds tailored, strategic, and role-specific.
11. Highlight concrete impact, measurable outcomes, and transferable strengths that map directly to the job.
12. Use a clear career progression flow: where I started, how I grew, key transitions, and what led me to apply for this role now.
13. Prioritize keywords and responsibilities from the job description when phrasing the narrative.
14. Do not invent experience not present in the resume; if details are missing, stay high-confidence and realistic.
15. Clearly explain what motivated me to apply for this role and why this role is the right next step.
16. End with a concise closing statement that reinforces fit and enthusiasm for this role.
17. Avoid technical jargon, acronyms, and buzzwords; use plain, human language that any interviewer can follow.
18. If a technical term is unavoidable, explain it in simple everyday wording.
19. Do not use section headers, titles, labels, bullet points, or numbered lists.
20. Do not include conversational opening pleasantries (for example: "Hi", "Thanks for having me", "Great to meet you").
21. Return the response as a conversation-like personal story about me in paragraph form only (2-3 cohesive paragraphs).
22. Keep the voice natural and spoken, as if I am talking in a general professional conversation.
23. Start naturally with a spoken opener in this style: "Well, I've been ..." and then explain how my experience built over time.
24. Emphasize progression over time, with examples of what I learned and how that led me to apply for this role.
25. Do not include title-like starters such as "Tell me about yourself:", "Background:", "Who I am:", or similar label text.
26. Do not use words like "excited", "thrilled", "glad", or "motivated".
27. Output only markdown text (no JSON, no HTML).
""".strip()


def build_elevator_pitch_prompt(job_role_details: str) -> str:
    return f"""
You are an expert interview coach.

Create a simple 30-second elevator pitch based on the role details below.

Job Role Details:
{job_role_details}

Instructions:
1. Write in first person as if I am speaking to an executive who asked what I do.
2. Keep it professional-conversational, clear, and natural, not robotic.
3. Keep it short enough for about 30 seconds at a medium pace.
4. Keep the total length between 50 and 75 words, and never exceed 75 words.
5. Focus on: who I am, what I do day-to-day, and the value I create.
6. Use plain language and avoid technical jargon or buzzwords.
7. Follow this structure every time:
    - Sentence 1: My current role context and how I contribute with my team.
    - Sentence 2: The specific business value I drive and the kinds of outcomes I create.
    - Sentence 3: Why this work matters to the business in simple terms.
8. Include a balance of people-oriented strengths and execution strengths (for example: collaboration, creative problem-solving, and data-informed decisions).
9. Keep the tone similar to this style: confident, direct, and grounded in impact, without sounding scripted.
10. Use short, clear sentences that sound like spoken conversation.
11. Prefer wording like "My role is...", "I work on...", "I also help...", and "That has helped...".
12. Avoid formal phrasing like "This enables me" or "propel the organization's success".
13. Keep language simple and natural, like I am speaking in a quick hallway conversation.
14. Avoid generic closing lines such as "has prepared me well for this role" or "my background in ... has prepared me".
15. Do not use titles, section headers, bullets, or numbered lists.
16. Return a single cohesive paragraph only.
17. Speak directly as if I am talking to one executive in front of me.
18. Do not include lead-in phrases like "Here is my narrative", "Here is my elevator pitch", "Sure", or "Absolutely".
19. Start immediately with the pitch content itself.
20. Do not use greeting/opening words such as "Hey", "Hi", or "Hello".
21. Do not use the word "excited".
22. Do not use phrases like "strong fit" or "perfect fit".
23. Output only markdown text (no JSON, no HTML).
""".strip()


def cap_narrative_to_medium_two_minutes(text: str, max_words: int = 260) -> str:
    words = text.split()
    if len(words) <= max_words:
        return text

    trimmed = " ".join(words[:max_words]).strip()
    last_boundary = max(trimmed.rfind("."), trimmed.rfind("!"), trimmed.rfind("?"))
    if last_boundary > int(len(trimmed) * 0.6):
        trimmed = trimmed[: last_boundary + 1]
    else:
        trimmed = trimmed.rstrip(" ,;:-") + "."

    return trimmed


def cleanup_narrative_format(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    cleaned_lines = []

    for line in lines:
        if not line:
            cleaned_lines.append("")
            continue

        if line.startswith("#"):
            continue
        lower = line.lower()
        if lower.startswith((
            "introduction:",
            "intro:",
            "background:",
            "closing:",
            "why i'm a fit:",
            "tell me about yourself:",
            "who i am:",
            "summary:",
        )):
            line = line.split(":", 1)[1].strip() if ":" in line else ""

        if line.endswith(":") and len(line.split()) <= 8:
            continue

        if line.startswith(("- ", "* ", "• ")):
            line = line[2:].strip()
        if len(line) > 2 and line[0].isdigit() and line[1] in ".)" and line[2] == " ":
            line = line[3:].strip()

        if line:
            cleaned_lines.append(line)

    text_flat = " ".join(part for part in cleaned_lines if part).strip()
    return text_flat


def strip_leading_preface(text: str) -> str:
    lower_text = text.lower().strip()
    preface_starts = [
        "here is my narrative",
        "here's my narrative",
        "here is my elevator pitch",
        "here's my elevator pitch",
        "here is your elevator pitch",
        "here's your elevator pitch",
        "this is my elevator pitch",
        "certainly",
        "absolutely",
        "sure",
    ]

    for phrase in preface_starts:
        if lower_text.startswith(phrase):
            for sep in [":", ".", "!", "?"]:
                idx = text.find(sep)
                if idx != -1 and idx < 140:
                    return text[idx + 1 :].strip()
            return ""

    return text


def soften_robotic_tone(text: str) -> str:
    replacements = {
        "I am ": "I'm ",
        "I have ": "I've ",
        "I would ": "I'd ",
        "do not": "don't",
        "cannot": "can't",
        "utilize": "use",
        "leverage": "use",
        "moreover": "also",
        "therefore": "so",
        "in conclusion": "overall",
        "I am confident": "I'm confident",
    }

    updated = text
    for old, new in replacements.items():
        updated = updated.replace(old, new)

    return updated


def enforce_elevator_word_rules(text: str) -> str:
    cleaned = text.strip()

    for greeting in ["Hey", "Hi", "Hello", "hey", "hi", "hello"]:
        if cleaned.startswith(greeting + " "):
            cleaned = cleaned[len(greeting) + 1 :].strip(" ,.!?-")
            break

    cleaned = cleaned.replace("excited", "motivated")
    cleaned = cleaned.replace("Excited", "Motivated")
    cleaned = cleaned.replace("strong fit", "good match")
    cleaned = cleaned.replace("Strong fit", "Good match")
    cleaned = cleaned.replace("perfect fit", "good match")
    cleaned = cleaned.replace("Perfect fit", "Good match")

    generic_phrases = [
        "has prepared me well for this role",
        "have prepared me well for this role",
        "my background in understanding customer behavior and making data-informed decisions",
    ]
    for phrase in generic_phrases:
        cleaned = cleaned.replace(phrase, "my hands-on work and results")
        cleaned = cleaned.replace(phrase.capitalize(), "My hands-on work and results")

    return cleaned


def enforce_story_word_rules(text: str) -> str:
    cleaned = text.strip()
    disallowed = ["excited", "Excited", "thrilled", "Thrilled", "glad", "Glad", "motivated", "Motivated"]
    for word in disallowed:
        cleaned = cleaned.replace(word, "")
    cleaned = " ".join(cleaned.split())
    return cleaned


def simplify_jargon(text: str) -> str:
    replacements = {
        "cross-functional": "across teams",
        "stakeholders": "the people involved",
        "end-to-end": "from start to finish",
        "strategic": "well-planned",
        "optimized": "improved",
        "optimization": "improvement",
        "synergy": "teamwork",
        "KPI": "key result",
        "KPIs": "key results",
        "scalable": "able to grow",
        "bandwidth": "time and capacity",
        "roadmap": "plan",
    }

    updated = text
    for old, new in replacements.items():
        updated = updated.replace(old, new)
        updated = updated.replace(old.title(), new.capitalize())

    return updated


def reset_all_fields():
    st.session_state.behavioral_question = ""
    st.session_state.situation_text = ""
    st.session_state.task_text = ""
    st.session_state.action_text = ""
    st.session_state.result_text = ""
    st.session_state.ai_followups = False
    st.session_state.final_answer = ""


def _render_body_html(lines: list[str], force_numbered: bool = False, allow_auto_numbered: bool = True) -> str:
    non_empty = [line for line in lines if line.strip()]
    is_numbered = non_empty and all(re.match(r"^\d+[\.)]\s+", line) for line in non_empty)
    if non_empty and ((allow_auto_numbered and is_numbered) or force_numbered):
        items = []
        for line in non_empty:
            item = re.sub(r"^\d+[\.)]\s+", "", line).strip()
            item = re.sub(r"^[-•]\s+", "", item).strip()
            items.append(f"<li>{html.escape(item)}</li>")
        return f'<ol class="star-list">{"".join(items)}</ol>'

    escaped = html.escape("\n".join(lines).strip())
    return escaped.replace("\n", "<br>")


def format_star_response_html(raw_text: str) -> str:
    cleaned = raw_text.replace("*", "").replace("\r\n", "\n")
    cleaned = re.sub(r"(?m)^\s*#{1,6}\s*", "", cleaned).strip()
    if not cleaned:
        return ""

    header_pattern = re.compile(
        r"^(?:The\s+)?(Situation|Task|Action(?:\s+Plan)?|Result|(?:Tailored|Customized|Additional)?\s*Follow[\s-]?up Questions?)"
        r"(?:\s*\([^)]*\))?(?:\s*[-\u2013\u2014]\s*[^:]*)?\s*:?\s*(.*)$",
        re.IGNORECASE,
    )
    lines = cleaned.split("\n")

    sections = []
    merged_sections: dict[str, list[str]] = {
        "Situation": [],
        "Task": [],
        "Action Plan": [],
        "Result": [],
        "Follow-up Questions": [],
    }
    current_header = None
    current_lines: list[str] = []
    orphan_lines: list[str] = []

    def _flush_section() -> None:
        nonlocal current_header, current_lines
        if not current_header:
            return
        is_followups = bool(re.search(r"Follow[\s-]?up Questions?", current_header, re.IGNORECASE))
        is_action = bool(re.match(r"^Action(?:\s+Plan)?$", current_header, re.IGNORECASE))
        section_lines = [
            re.sub(r"^\s*\d+[\.)]\s+", "", re.sub(r"^\s*[-•]\s+", "", line)).strip()
            for line in current_lines
        ]
        if is_followups:
            canonical_header = "Follow-up Questions"
        elif is_action:
            canonical_header = "Action Plan"
        elif re.match(r"^Situation$", current_header, re.IGNORECASE):
            canonical_header = "Situation"
        elif re.match(r"^Task$", current_header, re.IGNORECASE):
            canonical_header = "Task"
        elif re.match(r"^Result$", current_header, re.IGNORECASE):
            canonical_header = "Result"
        else:
            canonical_header = current_header

        if canonical_header in merged_sections:
            for line in section_lines:
                if line:
                    merged_sections[canonical_header].append(line)
        else:
            body_html = _render_body_html(
                section_lines,
                force_numbered=(is_followups or is_action),
                allow_auto_numbered=False,
            )
            sections.append(
                f'<div class="star-section"><div class="star-header">{html.escape(canonical_header)}</div>'
                f'<div class="star-body">{body_html}</div></div>'
            )
        current_header = None
        current_lines = []

    found_header = False
    for line in lines:
        match = header_pattern.match(line.strip())
        if match:
            found_header = True
            _flush_section()
            current_header = match.group(1)
            trailing = match.group(2).strip()
            current_lines = [trailing] if trailing else []
        else:
            if current_header:
                current_lines.append(line)
            else:
                orphan_lines.append(line)

    _flush_section()

    # Render known sections in fixed order, merging repeated headings.
    for key in ["Situation", "Task", "Action Plan", "Result", "Follow-up Questions"]:
        merged_lines = merged_sections[key]
        if not merged_lines:
            continue

        if key == "Action Plan":
            intro_lines: list[str] = []
            action_steps = merged_lines
            if merged_lines:
                first_line = merged_lines[0].strip()
                if re.search(r"following\s+steps\s*:?$", first_line, re.IGNORECASE):
                    intro_lines = [first_line]
                    action_steps = merged_lines[1:]

            intro_html = ""
            if intro_lines:
                intro_html = _render_body_html(intro_lines, force_numbered=False, allow_auto_numbered=False)
                intro_html = f"<div>{intro_html}</div>"

            steps_html = _render_body_html(action_steps, force_numbered=True, allow_auto_numbered=False)
            body_html = f"{intro_html}{steps_html}"
        else:
            body_html = _render_body_html(
                merged_lines,
                force_numbered=(key == "Follow-up Questions"),
                allow_auto_numbered=False,
            )
        if key == "Situation":
            display_header = "The Situation:"
        elif key == "Task":
            display_header = "The Task:"
        elif key == "Action Plan":
            display_header = "The Action Plan:"
        elif key == "Result":
            display_header = "The Result:"
        else:
            display_header = "Follow-up Questions"

        sections.append(
            f'<div class="star-section"><div class="star-header">{html.escape(display_header)}</div>'
            f'<div class="star-body">{body_html}</div></div>'
        )

    if found_header and sections:
        if any(line.strip() for line in orphan_lines):
            orphan_html = html.escape("\n".join(orphan_lines).strip()).replace("\n", "<br>")
            return f'<div class="star-section"><div class="star-body">{orphan_html}</div></div>' + "".join(sections)
        return "".join(sections)

    # Fallback: render plain non-bold text if sections aren't present.
    fallback = html.escape(cleaned).replace("\n", "<br>")
    return f'<div class="star-body">{fallback}</div>'


if "final_answer" not in st.session_state:
    st.session_state.final_answer = ""
if "behavioral_question" not in st.session_state:
    st.session_state.behavioral_question = ""
if "situation_text" not in st.session_state:
    st.session_state.situation_text = ""
if "task_text" not in st.session_state:
    st.session_state.task_text = ""
if "action_text" not in st.session_state:
    st.session_state.action_text = ""
if "result_text" not in st.session_state:
    st.session_state.result_text = ""
if "ai_followups" not in st.session_state:
    st.session_state.ai_followups = False

st.markdown(
    """
    <div class="brand-row">
        <div class="star-logo" aria-label="Star logo">&#9733;</div>
        <h1 class="brand-title">AI-STAR Interview Behavioral Question &amp; Narrative Builder</h1>
    </div>
    """,
    unsafe_allow_html=True,
)

tab_star, tab_story, tab_elevator = st.tabs(
    ["STAR Simulator", "Tell Me About Yourself", "Elevator Speech"]
)

with tab_star:
    question_left, question_right = st.columns([1, 1])
    with question_left:
        behavioral_question = st.text_input(
            "ENTER BEHAVIORAL QUESTION",
            key="behavioral_question",
            placeholder="Enter Question...",
        )
    with question_right:
        st.write("")

    left_col, right_col = st.columns([1, 1])

    with left_col:
        st.subheader("CRAFT YOUR STAR NARRATIVE")

        situation_text = st.text_area(
            "(S) SITUATION (Background & Context)",
            key="situation_text",
            placeholder="Enter situation details...",
            height=80,
        )
        task_text = st.text_area(
            "(T) TASK (The Responsibility & Goal)",
            key="task_text",
            placeholder="Enter task details...",
            height=80,
        )
        action_text = st.text_area(
            "(A) ACTION (Detailed Steps taken)",
            key="action_text",
            placeholder="Enter action details...",
            height=80,
        )
        result_text = st.text_area(
            "(R) RESULT (The Outcome & Impact)",
            key="result_text",
            placeholder="Enter result details...",
            height=80,
        )

        ai_followups = st.checkbox(
            "AI ENHANCEMENT: Get 2 customized AI follow-up questions based on my STAR narrative.",
            key="ai_followups",
        )

        st.button("Reset Fields", on_click=reset_all_fields)

        all_star_complete = all(
            [
                situation_text.strip(),
                task_text.strip(),
                action_text.strip(),
                result_text.strip(),
            ]
        )
        question_complete = bool(behavioral_question.strip())
        star_has_profanity = any_profanity_in_fields(
            [behavioral_question, situation_text, task_text, action_text, result_text]
        )

        if star_has_profanity:
            st.error("Profanity detected in the STAR prompts. Remove all profanity to enable generation.")

        if not all_star_complete or not question_complete:
            st.warning("Complete all 4 STAR fields and enter the behavioral question to enable AI generation.")

        generate_clicked = st.button(
            "Generate Your AI Narrative",
            disabled=not (all_star_complete and question_complete) or star_has_profanity,
        )

        if generate_clicked:
            if not all_star_complete:
                st.warning("Please fill in all STAR fields before generating.")
                st.stop()

            if not behavioral_question.strip():
                st.warning("Please enter the behavioral question before generating.")
                st.stop()

            if star_has_profanity:
                st.error("Profanity detected in the STAR prompts. Remove all profanity before generating.")
                st.stop()

            hf_token = get_hf_token()
            if not hf_token:
                st.error("Missing Hugging Face token. Add HF_TOKEN (preferred) or HF_API_TOKEN to environment variables/secrets.")
                st.stop()

            prompt = build_prompt(
                question=behavioral_question,
                situation=situation_text,
                task=task_text,
                action=action_text,
                result=result_text,
                add_followups=ai_followups,
            )

            with st.spinner("Generating your final STAR answer..."):
                try:
                    generated_text = call_hf_inference(prompt, hf_token)
                    st.session_state.final_answer = generated_text
                except Exception as gen_error:
                    st.error(f"Generation failed: {gen_error}")

    with right_col:
        st.subheader("YOUR STAR NARRATIVE")

        if st.session_state.final_answer:
            rendered_response_html = format_star_response_html(st.session_state.final_answer)
            st.markdown(f'<div class="answer-shell">{rendered_response_html}</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="answer-shell">&nbsp;</div>', unsafe_allow_html=True)

        if st.session_state.final_answer:
            docx_bytes = build_docx_bytes(
                question=behavioral_question,
                situation=situation_text,
                task=task_text,
                action=action_text,
                result=result_text,
                final_answer=st.session_state.final_answer,
            )

            st.download_button(
                "Download Narrative as MS Word Doc (.docx)",
                data=docx_bytes,
                file_name="ai_star_answer.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.download_button(
                "Download Narrative as MS Word Doc (.docx)",
                data=b"",
                file_name="ai_star_answer.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                disabled=True,
                use_container_width=True,
            )

with tab_story:
    st.subheader("Tell My Story Interview Simulator")
    story_left_col, story_right_col = st.columns([1.15, 1], gap="large")

    with story_left_col:
        st.markdown("### Enter Job Description")
        project_description = st.text_area(
            "Enter Job Description",
            key="project_description",
            height=220,
            placeholder="Paste or type the job description here...",
            label_visibility="collapsed",
        )
        story_has_profanity = has_profanity(project_description)

        if story_has_profanity:
            st.error("Profanity detected in the Tell Me About Yourself prompt. Remove it to enable generation.")

        if project_description.strip():
            st.markdown("### Upload Resume")
            uploaded_resume = st.file_uploader(
                "Upload your resume (PDF, DOCX, or TXT)",
                type=["pdf", "docx", "txt"],
                key="resume_upload",
            )

            if uploaded_resume is not None:
                st.success(f"Uploaded: {uploaded_resume.name}")
                if st.button("Remove uploaded resume", key="remove_resume"):
                    st.session_state.pop("resume_upload", None)
                    st.session_state.pop("resume_text", None)
                    st.rerun()
        else:
            st.info("Add your job description to unlock resume upload.")
            uploaded_resume = None

    with story_right_col:
        st.markdown("### Narrative Output")
        resume_uploaded = uploaded_resume is not None
        resume_text_for_gate = ""
        story_resume_has_profanity = False

        if resume_uploaded:
            resume_text_for_gate = extract_resume_text(uploaded_resume)
            st.session_state["resume_text"] = resume_text_for_gate
            story_resume_has_profanity = has_profanity(resume_text_for_gate)
            if story_resume_has_profanity:
                st.error("Profanity detected in uploaded resume content. Remove it to enable generation.")

        if resume_uploaded:
            generate_narrative = st.button(
                "Generate my Narrative",
                use_container_width=True,
                key="generate_narrative_btn",
                disabled=story_has_profanity or story_resume_has_profanity,
            )
        else:
            generate_narrative = False
            st.info("Upload your resume to unlock Generate my Narrative.")

        if generate_narrative:
            if not project_description.strip():
                st.error("Please enter a job description first.")
            elif story_has_profanity:
                st.error("Profanity detected. Remove all profanity before generating.")
            elif story_resume_has_profanity:
                st.error("Profanity detected in uploaded resume content. Remove all profanity before generating.")
            elif uploaded_resume is None:
                st.error("Please upload your resume.")
            else:
                resume_text = resume_text_for_gate
                st.session_state["resume_text"] = resume_text

                if not resume_text.strip():
                    st.error("The resume appears empty or unreadable. Try a different file.")
                else:
                    with st.spinner("Generating your interview narrative..."):
                        try:
                            hf_token = get_hf_token()
                            if not hf_token:
                                st.error(
                                    "Missing Hugging Face token. Add HF_TOKEN (preferred) or HF_API_TOKEN to environment variables/secrets."
                                )
                                st.stop()

                            prompt = build_narrative_prompt(project_description.strip(), resume_text)
                            narrative_md = call_hf_inference(prompt, hf_token, max_tokens=2000)
                            narrative_md = cleanup_narrative_format(narrative_md)
                            narrative_md = soften_robotic_tone(narrative_md)
                            narrative_md = simplify_jargon(narrative_md)
                            narrative_md = enforce_story_word_rules(narrative_md)
                            narrative_md = cap_narrative_to_medium_two_minutes(narrative_md, max_words=260)
                            st.session_state["narrative_md"] = narrative_md
                        except Exception as exc:
                            st.error(f"Hugging Face API error: {exc}")

        if "narrative_md" in st.session_state:
            st.markdown("---")
            st.markdown(st.session_state["narrative_md"])

with tab_elevator:
    st.subheader("Create Elevator Pitch")
    elevator_left_col, elevator_right_col = st.columns([1.15, 1], gap="large")

    with elevator_left_col:
        elevator_role_details = st.text_area(
            "Enter your Work Details as an example: 1. My Role with the Company is? 2. How does my work help the Team or add value to the Company. 3. What Problem do I help solve? 4. What Outcome do I help create?",
            key="elevator_role_details",
            height=220,
            placeholder=(
                "1) My Role with the Company is...\n"
                "2) My work helps the team/company by...\n"
                "3) A problem I help solve is...\n"
                "4) An outcome I help create is..."
            ),
            label_visibility="collapsed",
        )
        elevator_has_profanity = has_profanity(elevator_role_details)

        if elevator_has_profanity:
            st.error("Profanity detected in the Elevator Speech prompt. Remove it to enable generation.")

    with elevator_right_col:
        st.markdown("### 30-Second Elevator Pitch")
        submit_elevator = st.button(
            "Submit",
            use_container_width=True,
            key="submit_elevator_btn",
            disabled=elevator_has_profanity,
        )

        if submit_elevator:
            if not elevator_role_details.strip():
                st.error("Please enter your job role details first.")
            elif elevator_has_profanity:
                st.error("Profanity detected. Remove all profanity before generating.")
            else:
                with st.spinner("Generating your elevator pitch..."):
                    try:
                        hf_token = get_hf_token()
                        if not hf_token:
                            st.error("Missing Hugging Face token. Add HF_TOKEN (preferred) or HF_API_TOKEN to environment variables/secrets.")
                            st.stop()

                        elevator_prompt = build_elevator_pitch_prompt(elevator_role_details.strip())
                        elevator_md = call_hf_inference(elevator_prompt, hf_token, max_tokens=700)
                        elevator_md = cleanup_narrative_format(elevator_md)
                        elevator_md = strip_leading_preface(elevator_md)
                        elevator_md = soften_robotic_tone(elevator_md)
                        elevator_md = simplify_jargon(elevator_md)
                        elevator_md = enforce_elevator_word_rules(elevator_md)
                        elevator_md = cap_narrative_to_medium_two_minutes(elevator_md, max_words=85)
                        st.session_state["elevator_pitch_md"] = elevator_md
                    except Exception as exc:
                        st.error(f"Hugging Face API error: {exc}")

        if "elevator_pitch_md" in st.session_state:
            st.markdown("---")
            st.markdown(st.session_state["elevator_pitch_md"])
